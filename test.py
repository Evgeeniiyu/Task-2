import requests
from time import time
from docx import Document
import os
from dotenv import load_dotenv

load_dotenv()

BASE_URL = os.getenv("API_BASE_URL", "http://127.0.0.1:5000")

def generate_data(num_records):
    """Генерация данных для вставки"""
    return [{"name": f"Name {i}", "value": i} for i in range(num_records)]

def test_api_performance(num_records, results):
    """Тестирование производительности API"""
    print(f"\n--- Testing API with {num_records} records ---")

    response = requests.post(f"{BASE_URL}/setup")
    if response.status_code != 200:
        print("Error: Failed to setup the table")
        return

    # INSERT
    start = time()
    batch_size = 10000
    data = generate_data(num_records)
    for i in range(0, num_records, batch_size):
        response = requests.post(f"{BASE_URL}/insert", json={"data": data[i:i+batch_size]})
        if response.status_code != 200:
            print(f"Error: Failed to insert data at batch {i}")
            return
    insert_time = time() - start
    print(f"INSERT: {insert_time:.2f} sec")

    # SELECT
    start = time()
    response = requests.get(f"{BASE_URL}/select")
    if response.status_code != 200:
        print("Error: Failed to select data")
        return
    select_time = time() - start
    print(f"SELECT: {select_time:.2f} sec")

    # UPDATE
    start = time()
    response = requests.put(f"{BASE_URL}/update")
    if response.status_code != 200:
        print("Error: Failed to update data")
        return
    update_time = time() - start
    print(f"UPDATE: {update_time:.2f} sec")

    # DELETE
    start = time()
    response = requests.delete(f"{BASE_URL}/delete")
    if response.status_code != 200:
        print("Error: Failed to delete data")
        return
    delete_time = time() - start
    print(f"DELETE: {delete_time:.2f} sec")

    results.append([num_records, insert_time, select_time, update_time, delete_time])

def save_results_to_word(results):
    """Сохранение результатов в Word файл"""
    doc = Document()
    doc.add_heading("Результаты тестирования производительности API", level=1)

    # Создание таблицы
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Кол-во записей"
    hdr_cells[1].text = "INSERT (сек)"
    hdr_cells[2].text = "SELECT (сек)"
    hdr_cells[3].text = "UPDATE (сек)"
    hdr_cells[4].text = "DELETE (сек)"

    # Добавление данных
    for row in results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row[0])
        row_cells[1].text = f"{row[1]:.2f}"
        row_cells[2].text = f"{row[2]:.2f}"
        row_cells[3].text = f"{row[3]:.2f}"
        row_cells[4].text = f"{row[4]:.2f}"

    # Сохранение файла
    file_name = "Test.docx"
    doc.save(file_name)
    print(f"\nРезультаты сохранены в файл '{file_name}'.")

results = []
for records in [1000, 10000, 100000, 1000000]:
    test_api_performance(records, results)

save_results_to_word(results)
